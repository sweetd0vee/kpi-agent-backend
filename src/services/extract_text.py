"""
Извлечение текста из загруженных файлов (PDF, DOCX, XLSX, TXT).
Для последующей передачи в LLM и преобразования в JSON.
"""
import io
from pathlib import Path

from pypdf import PdfReader
from docx import Document as DocxDocument
import openpyxl


def extract_text_from_file(file_path: Path, content_type: str | None, filename: str) -> str:
    """
    Извлекает текст из файла по расширению или content_type.
    Поддерживаются: .txt, .pdf, .docx, .xlsx.
    """
    suffix = (file_path.suffix or "").lower()
    if content_type:
        ct = content_type.lower()
        if "pdf" in ct:
            suffix = ".pdf"
        elif "word" in ct or "docx" in ct:
            suffix = ".docx"
        elif "sheet" in ct or "excel" in ct or "xlsx" in ct:
            suffix = ".xlsx"
        elif "text/plain" in ct:
            suffix = ".txt"

    if suffix == ".txt":
        return file_path.read_text(encoding="utf-8", errors="replace")

    if suffix == ".pdf":
        reader = PdfReader(str(file_path))
        parts = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                parts.append("")
        return "\n".join(parts)

    if suffix == ".docx":
        doc = DocxDocument(str(file_path))
        return "\n".join(p.text for p in doc.paragraphs)

    if suffix in (".xlsx", ".xls"):
        wb = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
        parts = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                line = "\t".join(str(c) if c is not None else "" for c in row)
                if line.strip():
                    parts.append(line)
        wb.close()
        return "\n".join(parts)

    # Fallback: try read as text
    try:
        return file_path.read_text(encoding="utf-8", errors="replace")
    except Exception:
        return ""


def extract_text_from_bytes(data: bytes, filename: str, content_type: str | None = None) -> str:
    """Извлекает текст из байтов (например, из UploadFile)."""
    path = Path(filename)
    suffix = path.suffix.lower()
    if content_type:
        ct = content_type.lower()
        if "pdf" in ct:
            suffix = ".pdf"
        elif "word" in ct or "docx" in ct:
            suffix = ".docx"
        elif "sheet" in ct or "excel" in ct or "xlsx" in ct:
            suffix = ".xlsx"
        elif "text/plain" in ct:
            suffix = ".txt"

    if suffix == ".txt":
        return data.decode("utf-8", errors="replace")

    if suffix == ".pdf":
        reader = PdfReader(io.BytesIO(data))
        parts = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                parts.append("")
        return "\n".join(parts)

    if suffix == ".docx":
        doc = DocxDocument(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs)

    if suffix in (".xlsx", ".xls"):
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                line = "\t".join(str(c) if c is not None else "" for c in row)
                if line.strip():
                    parts.append(line)
        wb.close()
        return "\n".join(parts)

    try:
        return data.decode("utf-8", errors="replace")
    except Exception:
        return ""
